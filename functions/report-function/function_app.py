import azure.functions as func
import json
import os
import pyodbc
import tempfile
from datetime import datetime

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from azure.storage.blob import BlobServiceClient
from openai import AzureOpenAI


app = func.FunctionApp(http_auth_level=func.AuthLevel.FUNCTION)


# ============================================================
# 공통 유틸
# ============================================================

def get_sql_connection():
    conn_str = os.environ["SQL_CONNECTION_STRING"]
    return pyodbc.connect(conn_str)


def get_env_value(*names, default=None):
    for name in names:
        value = os.environ.get(name)
        if value:
            return value
    return default


def safe_json_loads(text):
    if not text:
        return {}

    cleaned = text.strip()

    if cleaned.startswith("```json"):
        cleaned = cleaned.replace("```json", "", 1).strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.replace("```", "", 1).strip()
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3].strip()

    return json.loads(cleaned)


def parse_json_text(text):
    if not text:
        return {}

    try:
        return json.loads(text)
    except Exception:
        try:
            return safe_json_loads(text)
        except Exception:
            return {}


def format_value(value, suffix=""):
    if value is None:
        return "-"
    return f"{value}{suffix}"


def safe_int(value, default=None):
    if value is None:
        return default
    try:
        return int(value)
    except Exception:
        return default


def safe_float(value, default=None):
    if value is None:
        return default
    try:
        return float(value)
    except Exception:
        return default


def get_dict_value(data, key, default=None):
    if isinstance(data, dict):
        return data.get(key, default)
    return default


def get_first_value(*values):
    for value in values:
        if value is not None:
            return value
    return None


def row_value(row, attr, default=None):
    if not row:
        return default
    try:
        return getattr(row, attr)
    except Exception:
        return default


def get_table_columns(conn, table_name, schema_name="dbo"):
    cursor = conn.cursor()
    cursor.execute("""
        SELECT COLUMN_NAME
        FROM INFORMATION_SCHEMA.COLUMNS
        WHERE TABLE_SCHEMA = ?
          AND TABLE_NAME = ?
    """, schema_name, table_name)

    return {row.COLUMN_NAME for row in cursor.fetchall()}


def build_optional_column_select(existing_columns, column_name):
    if column_name in existing_columns:
        return column_name
    return f"CAST(NULL AS FLOAT) AS {column_name}"


# ============================================================
# DOCX 한글 폰트 처리
# ============================================================

def set_rfonts(rpr, font_name="Malgun Gothic"):
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)

    rfonts.set(qn("w:ascii"), font_name)
    rfonts.set(qn("w:hAnsi"), font_name)
    rfonts.set(qn("w:eastAsia"), font_name)
    rfonts.set(qn("w:cs"), font_name)


def set_run_font(run, font_name="Malgun Gothic"):
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    set_rfonts(rpr, font_name)


def set_paragraph_font(paragraph, font_name="Malgun Gothic"):
    for run in paragraph.runs:
        set_run_font(run, font_name)


def set_style_font(doc, style_name, font_name="Malgun Gothic"):
    try:
        style = doc.styles[style_name]
        style.font.name = font_name
        rpr = style._element.get_or_add_rPr()
        set_rfonts(rpr, font_name)
    except Exception:
        pass


def apply_doc_font(doc, font_name="Malgun Gothic"):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]:
        set_style_font(doc, style_name, font_name)


def add_fonted_heading(doc, text, level=1, font_name="Malgun Gothic"):
    paragraph = doc.add_heading(text, level=level)
    set_paragraph_font(paragraph, font_name)
    return paragraph


def add_fonted_paragraph(doc, text, font_name="Malgun Gothic"):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text if text is not None else "")
    set_run_font(run, font_name)
    return paragraph


def set_cell_text(cell, text, font_name="Malgun Gothic"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text) if text is not None else "")
    set_run_font(run, font_name)


# ============================================================
# SQL 데이터 조회
# ============================================================

def fetch_market_feature(conn, interval_type, order_timestamp):
    cursor = conn.cursor()

    existing_columns = get_table_columns(conn, "market_feature")

    avg_spread_select = build_optional_column_select(existing_columns, "avg_spread")
    avg_bid_depth_select = build_optional_column_select(existing_columns, "avg_bid_depth")
    avg_ask_depth_select = build_optional_column_select(existing_columns, "avg_ask_depth")
    depth_imbalance_select = build_optional_column_select(existing_columns, "depth_imbalance")
    lp_active_ratio_select = build_optional_column_select(existing_columns, "lp_active_ratio")

    query = f"""
        SELECT TOP 1
            feature_id,
            window_timestamp,
            interval_type,
            return_rate,
            volatility,
            buy_order_count,
            sell_order_count,
            order_imbalance,
            sync_index,
            avg_slippage_rate,
            lp_trade_count,
            volume,
            {avg_spread_select},
            {avg_bid_depth_select},
            {avg_ask_depth_select},
            {depth_imbalance_select},
            {lp_active_ratio_select}
        FROM dbo.market_feature
        WHERE interval_type = ?
          AND window_timestamp <= ?
        ORDER BY window_timestamp DESC
    """

    cursor.execute(query, interval_type, order_timestamp)
    row = cursor.fetchone()

    # 주문 시각 이전 데이터가 없으면, 동일 interval_type의 최신 market_feature로 보완
    # Databricks 분석 산출물 저장 시각과 주문 로그 시각의 기준이 다르면
    # 기존 조건(window_timestamp <= order_timestamp)으로는 feature가 비어 표가 '-'로 나올 수 있음
    if row:
        return row

    fallback_query = f"""
        SELECT TOP 1
            feature_id,
            window_timestamp,
            interval_type,
            return_rate,
            volatility,
            buy_order_count,
            sell_order_count,
            order_imbalance,
            sync_index,
            avg_slippage_rate,
            lp_trade_count,
            volume,
            {avg_spread_select},
            {avg_bid_depth_select},
            {avg_ask_depth_select},
            {depth_imbalance_select},
            {lp_active_ratio_select}
        FROM dbo.market_feature
        WHERE interval_type = ?
        ORDER BY window_timestamp DESC
    """

    cursor.execute(fallback_query, interval_type)
    return cursor.fetchone()


def fetch_report_data(conn, agent_id, target_order_id, interval_type="1m"):
    cursor = conn.cursor()

    # 1. 주문 정보
    cursor.execute("""
        SELECT 
            order_id,
            agent_id,
            order_side,
            order_type,
            price,
            quantity,
            status,
            reject_reason,
            created_at
        FROM dbo.system_order
        WHERE agent_id = ? AND order_id = ?
    """, agent_id, target_order_id)
    order = cursor.fetchone()

    if not order:
        raise Exception("분석 대상 주문을 찾을 수 없습니다.")

    order_timestamp = order.created_at

    if not order_timestamp:
        raise Exception("분석 대상 주문의 created_at 값이 없습니다.")

    # 2. 체결 정보
    cursor.execute("""
        SELECT 
            SUM(CAST(price AS BIGINT) * quantity) * 1.0 / NULLIF(SUM(quantity), 0) AS average_execution_price,
            SUM(quantity) AS executed_quantity,
            SUM(CAST(price AS BIGINT) * quantity) AS trade_amount,
            COUNT(*) AS trade_count,
            SUM(CASE WHEN is_lp_trade = 1 THEN 1 ELSE 0 END) AS lp_trade_count
        FROM dbo.system_trade
        WHERE buy_order_id = ? OR sell_order_id = ?
    """, target_order_id, target_order_id)
    execution = cursor.fetchone()

    # 3. 주문 시점 기준 가장 가까운 이전 1초 OHLCV
    cursor.execute("""
        SELECT TOP 1
            candle_timestamp,
            close_price,
            volume,
            trade_count
        FROM dbo.market_ohlcv
        WHERE candle_timestamp <= ?
        ORDER BY candle_timestamp DESC
    """, order_timestamp)
    ohlcv = cursor.fetchone()

    # 4. 주문 시점 기준 가장 가까운 이전 기간별 OHLCV
    cursor.execute("""
        SELECT TOP 1
            candle_timestamp,
            interval_type,
            open_price,
            high_price,
            low_price,
            close_price,
            volume,
            trade_count
        FROM dbo.market_ohlcv_agg
        WHERE interval_type = ?
          AND candle_timestamp <= ?
        ORDER BY candle_timestamp DESC
    """, interval_type, order_timestamp)
    ohlcv_agg = cursor.fetchone()

    # 5. 주문 시점 기준 가장 가까운 이전 market_feature
    feature = fetch_market_feature(conn, interval_type, order_timestamp)

    # 6. 계좌 상태
    cursor.execute("""
        SELECT
            agent_id,
            strategy_type,
            deposit,
            balance,
            avg_buy_price,
            locked_deposit,
            locked_balance,
            realized_pnl,
            updated_at
        FROM dbo.account
        WHERE agent_id = ?
    """, agent_id)
    account = cursor.fetchone()

    # 7. 에이전트 최종 성과
    cursor.execute("""
        SELECT TOP 1
            performance_id,
            agent_id,
            strategy_type,
            final_deposit,
            final_balance,
            final_avg_buy_price,
            realized_pnl,
            unrealized_pnl,
            total_asset_value,
            return_rate,
            trade_count,
            rank_no,
            created_at
        FROM dbo.agent_performance
        WHERE agent_id = ?
        ORDER BY created_at DESC
    """, agent_id)
    performance = cursor.fetchone()

    # 8. 행동 분석 결과
    cursor.execute("""
        SELECT TOP 1
            analysis_id,
            agent_id,
            target_order_id,
            behavior_type,
            behavior_summary,
            input_feature_json,
            model_output_json,
            model_name,
            model_version,
            created_at
        FROM dbo.behavior_analysis_result
        WHERE agent_id = ? AND target_order_id = ?
        ORDER BY created_at DESC
    """, agent_id, target_order_id)
    behavior = cursor.fetchone()

    return {
        "order": order,
        "order_timestamp": order_timestamp,
        "execution": execution,
        "ohlcv": ohlcv,
        "ohlcv_agg": ohlcv_agg,
        "feature": feature,
        "account": account,
        "performance": performance,
        "behavior": behavior
    }


# ============================================================
# OpenAI 입력 데이터 구성
# ============================================================

def build_openai_input(agent_id, target_order_id, req_body, data):
    order = data["order"]
    order_timestamp = data["order_timestamp"]
    execution = data["execution"]
    ohlcv = data["ohlcv"]
    ohlcv_agg = data["ohlcv_agg"]
    feature = data["feature"]
    account = data["account"]
    performance = data["performance"]
    behavior = data["behavior"]

    interval_type = req_body.get("interval_type", "1m")

    behavior_input_features = {}
    behavior_model_output = {}

    if behavior:
        behavior_input_features = parse_json_text(behavior.input_feature_json)
        behavior_model_output = parse_json_text(behavior.model_output_json)

    if not isinstance(behavior_input_features, dict):
        behavior_input_features = {}

    if not isinstance(behavior_model_output, dict):
        behavior_model_output = {}

    # Databricks 신규 JSON 구조 대응
    # 실제 input_feature_json 최상위 구조:
    # report_request, selected_execution_context, last_n_execution_context,
    # all_trades_execution_context, execution, market_context, agent_cumulative_context
    report_request_json = get_dict_value(behavior_input_features, "report_request", {})
    execution_json = get_dict_value(behavior_input_features, "execution", {})
    market_context_json = get_dict_value(behavior_input_features, "market_context", {})
    agent_context_json = get_dict_value(behavior_input_features, "agent_cumulative_context", {})
    selected_execution_context_json = get_dict_value(behavior_input_features, "selected_execution_context", {})
    last_n_execution_context_json = get_dict_value(behavior_input_features, "last_n_execution_context", {})
    all_trades_execution_context_json = get_dict_value(behavior_input_features, "all_trades_execution_context", {})

    if not isinstance(report_request_json, dict):
        report_request_json = {}
    if not isinstance(execution_json, dict):
        execution_json = {}
    if not isinstance(market_context_json, dict):
        market_context_json = {}
    if not isinstance(agent_context_json, dict):
        agent_context_json = {}
    if not isinstance(selected_execution_context_json, dict):
        selected_execution_context_json = {}
    if not isinstance(last_n_execution_context_json, dict):
        last_n_execution_context_json = {}
    if not isinstance(all_trades_execution_context_json, dict):
        all_trades_execution_context_json = {}

    # 보고서 범위 선택: 기본은 Databricks가 지정한 LAST_N_TRADES 사용
    requested_scope_type = req_body.get(
        "scope_type",
        report_request_json.get("scope_type", "LAST_N_TRADES")
    )

    if requested_scope_type == "ALL_TRADES":
        selected_scope_execution_json = all_trades_execution_context_json
    elif requested_scope_type == "SELECTED_TRADES":
        selected_scope_execution_json = selected_execution_context_json
    else:
        selected_scope_execution_json = last_n_execution_context_json or selected_execution_context_json

    # current_price는 주문 시점 기준 가격으로 사용
    # Databricks 필수 저장 대상이 아니므로 Function에서 OHLCV 기준 조회
    current_price = None
    current_price_source = None

    if ohlcv and ohlcv.close_price is not None:
        current_price = int(ohlcv.close_price)
        current_price_source = "market_ohlcv_1s"
    elif ohlcv_agg and ohlcv_agg.close_price is not None:
        current_price = int(ohlcv_agg.close_price)
        current_price_source = "market_ohlcv_agg"

    # trade_count는 market_feature가 아니라 OHLCV 기준으로 사용
    market_trade_count = None
    if interval_type == "1s":
        market_trade_count = safe_int(row_value(ohlcv, "trade_count"))
    else:
        market_trade_count = safe_int(row_value(ohlcv_agg, "trade_count"))

    # Databricks execution JSON 우선, 없으면 system_trade 집계값 보완
    average_execution_price = get_first_value(
        safe_float(execution_json.get("average_execution_price")),
        safe_float(execution_json.get("avg_execution_price")),
        safe_float(row_value(execution, "average_execution_price"))
    )

    executed_quantity = get_first_value(
        safe_int(execution_json.get("executed_quantity")),
        safe_int(execution_json.get("filled_quantity")),
        safe_int(row_value(execution, "executed_quantity")),
        0
    )

    trade_amount = get_first_value(
        safe_int(execution_json.get("executed_amount")),
        safe_int(execution_json.get("execution_amount")),
        safe_int(row_value(execution, "trade_amount")),
        0
    )

    execution_trade_count = get_first_value(
        safe_int(execution_json.get("execution_trade_count")),
        safe_int(row_value(execution, "trade_count")),
        0
    )

    lp_trade_count = get_first_value(
        safe_int(execution_json.get("lp_trade_count")),
        safe_int(selected_scope_execution_json.get("lp_trade_count")),
        safe_int(row_value(execution, "lp_trade_count")),
        0
    )

    # 주문별 슬리피지율 우선순위
    # 1. input_feature_json.execution.slippage_rate
    # 2. 기존 flat target_order_slippage_rate
    # 3. market_context.avg_slippage_rate
    # 4. market_feature.avg_slippage_rate
    slippage_rate = get_first_value(
        safe_float(execution_json.get("slippage_rate")),
        safe_float(behavior_input_features.get("target_order_slippage_rate")),
        safe_float(market_context_json.get("avg_slippage_rate")),
        safe_float(market_context_json.get("avg_market_slippage_rate")),
        safe_float(row_value(feature, "avg_slippage_rate"))
    )

    # 예상 기회손익: Databricks execution JSON 우선, 없으면 Function에서 보완 계산
    opportunity_loss = safe_int(execution_json.get("opportunity_loss"))

    if opportunity_loss is None and current_price is not None and average_execution_price is not None and executed_quantity:
        if order.order_side == "BUY":
            opportunity_loss = int((average_execution_price - current_price) * executed_quantity)
        elif order.order_side == "SELL":
            opportunity_loss = int((current_price - average_execution_price) * executed_quantity)

    order_summary_text = f"{order.order_type} {order.order_side} / {int(order.quantity)}주"

    main_causes = behavior_model_output.get("main_causes", [])
    if not isinstance(main_causes, list):
        main_causes = []

    feature_interval_type = (
        feature.interval_type
        if feature and feature.interval_type
        else interval_type
    )

    # 주문 행동 지표: agent_cumulative_context 우선, 없으면 기존 flat 구조 보완
    total_order_count = get_first_value(
        agent_context_json.get("total_order_count"),
        behavior_input_features.get("total_order_count")
    )

    market_order_count = get_first_value(
        agent_context_json.get("market_order_count"),
        behavior_input_features.get("market_order_count")
    )

    sell_order_count_from_behavior = get_first_value(
        agent_context_json.get("sell_order_count"),
        behavior_input_features.get("sell_order_count")
    )

    market_order_ratio = get_first_value(
        agent_context_json.get("market_order_ratio"),
        behavior_input_features.get("market_order_ratio")
    )

    sell_order_ratio = get_first_value(
        agent_context_json.get("sell_order_ratio"),
        behavior_input_features.get("sell_order_ratio")
    )

    if sell_order_ratio is None:
        total_order_count_for_ratio = safe_float(total_order_count)
        sell_order_count_for_ratio = safe_float(sell_order_count_from_behavior)
        if total_order_count_for_ratio and sell_order_count_for_ratio is not None:
            sell_order_ratio = sell_order_count_for_ratio / total_order_count_for_ratio

    # 기존 avg_order_quantity가 없으면 선택된 체결 범위의 평균 체결 수량으로 보완
    avg_order_quantity = get_first_value(
        agent_context_json.get("avg_order_quantity"),
        behavior_input_features.get("avg_order_quantity")
    )

    if avg_order_quantity is None:
        scope_trade_count_for_avg = safe_float(selected_scope_execution_json.get("trade_count"))
        scope_total_quantity_for_avg = safe_float(selected_scope_execution_json.get("total_quantity"))
        if scope_trade_count_for_avg and scope_total_quantity_for_avg is not None:
            avg_order_quantity = scope_total_quantity_for_avg / scope_trade_count_for_avg

    limit_order_count = get_first_value(
        agent_context_json.get("limit_order_count"),
        behavior_input_features.get("limit_order_count")
    )

    # 시장 지표: market_context 우선, 없으면 market_feature 보완
    market_return_rate = get_first_value(
        safe_float(market_context_json.get("return_rate")),
        safe_float(market_context_json.get("avg_market_return_rate")),
        safe_float(row_value(feature, "return_rate"))
    )

    market_volatility = get_first_value(
        safe_float(market_context_json.get("volatility")),
        safe_float(market_context_json.get("avg_market_volatility")),
        safe_float(row_value(feature, "volatility"))
    )

    market_order_imbalance = get_first_value(
        safe_float(market_context_json.get("order_imbalance")),
        safe_float(agent_context_json.get("order_imbalance")),
        safe_float(row_value(feature, "order_imbalance"))
    )

    market_sync_index = get_first_value(
        safe_float(market_context_json.get("sync_index")),
        safe_float(market_context_json.get("avg_sync_index")),
        safe_float(market_context_json.get("market_sync_index")),
        safe_float(market_context_json.get("avg_market_sync_index")),
        safe_float(behavior_input_features.get("sync_index")),
        safe_float(row_value(feature, "sync_index"))
    )

    market_avg_slippage_rate = get_first_value(
        safe_float(market_context_json.get("avg_slippage_rate")),
        safe_float(market_context_json.get("avg_market_slippage_rate")),
        safe_float(row_value(feature, "avg_slippage_rate"))
    )

    # 보고서 표의 구간 거래량은 분석 범위 체결 수량을 우선 사용
    # market_window_count는 feature row 개수이므로 거래량 의미와 달라 후순위로 둠
    market_volume = get_first_value(
        safe_int(selected_scope_execution_json.get("total_quantity")),
        safe_int(all_trades_execution_context_json.get("total_quantity")),
        safe_int(market_context_json.get("volume")),
        safe_int(row_value(feature, "volume")),
        safe_int(market_context_json.get("market_window_count"))
    )

    market_context_trade_count = get_first_value(
        safe_int(market_context_json.get("trade_count")),
        safe_int(selected_scope_execution_json.get("trade_count")),
        safe_int(selected_scope_execution_json.get("actual_trade_count")),
        market_trade_count
    )

    avg_spread = get_first_value(
        safe_float(market_context_json.get("avg_spread")),
        safe_float(market_context_json.get("avg_market_spread")),
        safe_float(row_value(feature, "avg_spread"))
    )

    avg_bid_depth = get_first_value(
        safe_float(market_context_json.get("avg_bid_depth")),
        safe_float(market_context_json.get("avg_market_bid_depth")),
        safe_float(row_value(feature, "avg_bid_depth"))
    )

    avg_ask_depth = get_first_value(
        safe_float(market_context_json.get("avg_ask_depth")),
        safe_float(market_context_json.get("avg_market_ask_depth")),
        safe_float(row_value(feature, "avg_ask_depth"))
    )

    depth_imbalance = get_first_value(
        safe_float(market_context_json.get("depth_imbalance")),
        safe_float(market_context_json.get("avg_depth_imbalance")),
        safe_float(market_context_json.get("market_depth_imbalance")),
        safe_float(row_value(feature, "depth_imbalance"))
    )

    # JSON에 호가 깊이 불균형이 직접 없으면 평균 매수/매도호가 깊이로 계산
    if depth_imbalance is None:
        bid_depth_for_imbalance = safe_float(avg_bid_depth)
        ask_depth_for_imbalance = safe_float(avg_ask_depth)
        if bid_depth_for_imbalance is not None and ask_depth_for_imbalance is not None:
            depth_sum = bid_depth_for_imbalance + ask_depth_for_imbalance
            if depth_sum != 0:
                depth_imbalance = (bid_depth_for_imbalance - ask_depth_for_imbalance) / depth_sum

    lp_active_ratio = get_first_value(
        safe_float(market_context_json.get("lp_active_ratio")),
        safe_float(market_context_json.get("avg_lp_active_ratio")),
        safe_float(row_value(feature, "lp_active_ratio"))
    )

    return {
        "report_context": {
            "agent_id": agent_id,
            "strategy_type": (
                performance.strategy_type
                if performance and performance.strategy_type
                else account.strategy_type if account else None
            ),
            "symbol_name": req_body.get("symbol_name", "삼성전자 디지털 트윈 마켓"),
            "digital_twin_market_no": req_body.get("digital_twin_market_no", "DTM_001"),
            "analysis_start_date": get_first_value(
                req_body.get("analysis_start_date"),
                market_context_json.get("analysis_start_ts")
            ),
            "analysis_end_date": get_first_value(
                req_body.get("analysis_end_date"),
                market_context_json.get("analysis_end_ts"),
                market_context_json.get("analysis_cutoff_ts"),
                report_request_json.get("analysis_cutoff_ts")
            ),
            "report_type": req_body.get("report_type", "INVESTMENT_BEHAVIOR_DIAGNOSIS"),
            "target_order_timestamp": str(order_timestamp)
        },
        "performance_summary": {
            "final_deposit": int(performance.final_deposit) if performance and performance.final_deposit is not None else int(account.deposit) if account and account.deposit is not None else None,
            "final_balance": int(performance.final_balance) if performance and performance.final_balance is not None else int(account.balance) if account and account.balance is not None else None,
            "final_avg_buy_price": int(performance.final_avg_buy_price) if performance and performance.final_avg_buy_price is not None else int(account.avg_buy_price) if account and account.avg_buy_price is not None else None,
            "realized_pnl": int(performance.realized_pnl) if performance and performance.realized_pnl is not None else int(account.realized_pnl) if account and account.realized_pnl is not None else safe_int(agent_context_json.get("realized_pnl")),
            "unrealized_pnl": int(performance.unrealized_pnl) if performance and performance.unrealized_pnl is not None else safe_int(agent_context_json.get("unrealized_pnl")),
            "total_asset_value": int(performance.total_asset_value) if performance and performance.total_asset_value is not None else safe_int(agent_context_json.get("total_asset_value")),
            "return_rate": float(performance.return_rate) if performance and performance.return_rate is not None else safe_float(agent_context_json.get("return_rate")),
            "trade_count": int(performance.trade_count) if performance and performance.trade_count is not None else get_first_value(safe_int(agent_context_json.get("trade_count")), execution_trade_count),
            "rank_no": int(performance.rank_no) if performance and performance.rank_no is not None else safe_int(agent_context_json.get("rank_no"))
        },
        "order_summary": {
            "target_order_id": int(target_order_id),
            "target_order_timestamp": str(order_timestamp),
            "order_summary_text": order_summary_text,
            "order_side": order.order_side,
            "order_type": order.order_type,
            "order_price": int(order.price) if order.price is not None else None,
            "order_quantity": int(order.quantity),
            "status": order.status,
            "reject_reason": order.reject_reason,

            "total_order_count": total_order_count,
            "buy_order_count": int(feature.buy_order_count) if feature and feature.buy_order_count is not None else behavior_input_features.get("buy_order_count"),
            "sell_order_count": int(feature.sell_order_count) if feature and feature.sell_order_count is not None else sell_order_count_from_behavior,
            "market_order_count": market_order_count,
            "limit_order_count": limit_order_count,
            "market_order_ratio": market_order_ratio,
            "sell_order_ratio": sell_order_ratio,
            "avg_order_quantity": avg_order_quantity
        },
        "execution_summary": {
            "current_price": current_price,
            "current_price_source": current_price_source,
            "average_execution_price": average_execution_price,
            "executed_quantity": executed_quantity,
            "trade_amount": trade_amount,
            "trade_count": execution_trade_count,
            "slippage_rate": slippage_rate,
            "opportunity_loss": opportunity_loss,
            "lp_trade_count": lp_trade_count
        },
        "market_summary": {
            "interval_type": feature_interval_type,
            "current_price": current_price,
            "current_price_source": current_price_source,
            "ohlcv_1s_timestamp": str(ohlcv.candle_timestamp) if ohlcv else None,
            "ohlcv_agg_timestamp": str(ohlcv_agg.candle_timestamp) if ohlcv_agg else None,
            "feature_window_timestamp": get_first_value(
                str(feature.window_timestamp) if feature else None,
                market_context_json.get("analysis_cutoff_ts"),
                market_context_json.get("analysis_end_ts")
            ),

            "return_rate": market_return_rate,
            "volatility": market_volatility,
            "buy_order_count": int(feature.buy_order_count) if feature and feature.buy_order_count is not None else None,
            "sell_order_count": int(feature.sell_order_count) if feature and feature.sell_order_count is not None else None,
            "order_imbalance": market_order_imbalance,
            "sync_index": market_sync_index,
            "avg_slippage_rate": market_avg_slippage_rate,
            "lp_trade_count": int(feature.lp_trade_count) if feature and feature.lp_trade_count is not None else lp_trade_count,
            "volume": market_volume,
            "trade_count": market_context_trade_count,

            "avg_spread": avg_spread,
            "avg_bid_depth": avg_bid_depth,
            "avg_ask_depth": avg_ask_depth,
            "depth_imbalance": depth_imbalance,
            "lp_active_ratio": lp_active_ratio
        },
        "execution_context_summary": {
            "scope_type": requested_scope_type,
            "requested_trade_count": report_request_json.get("requested_trade_count"),
            "actual_trade_count": selected_scope_execution_json.get("actual_trade_count"),
            "selection_rule": selected_scope_execution_json.get("selection_rule"),
            "trade_count": selected_scope_execution_json.get("trade_count"),
            "buy_trade_count": selected_scope_execution_json.get("buy_trade_count"),
            "sell_trade_count": selected_scope_execution_json.get("sell_trade_count"),
            "total_quantity": selected_scope_execution_json.get("total_quantity"),
            "gross_trade_amount": selected_scope_execution_json.get("gross_trade_amount"),
            "net_quantity": selected_scope_execution_json.get("net_quantity"),
            "net_trade_amount": selected_scope_execution_json.get("net_trade_amount"),
            "weighted_avg_execution_price": selected_scope_execution_json.get("weighted_avg_execution_price"),
            "min_execution_price": selected_scope_execution_json.get("min_execution_price"),
            "max_execution_price": selected_scope_execution_json.get("max_execution_price"),
            "first_created_at": selected_scope_execution_json.get("first_created_at"),
            "last_created_at": selected_scope_execution_json.get("last_created_at")
        },
        "behavior_summary": {
            "behavior_type": behavior.behavior_type if behavior else agent_context_json.get("behavior_type"),
            "behavior_summary_text": behavior.behavior_summary if behavior else agent_context_json.get("behavior_summary"),
            "diagnosis_seed_text": agent_context_json.get("diagnosis_seed_text"),
            "prescription_seed_text": agent_context_json.get("prescription_seed_text"),
            "main_causes": main_causes,
            "input_features": behavior_input_features
        }
    }


# ============================================================
# Azure OpenAI 호출
# ============================================================

def call_openai(openai_input):
    endpoint = os.environ["AZURE_OPENAI_ENDPOINT"]
    api_key = os.environ["AZURE_OPENAI_API_KEY"]

    deployment_name = get_env_value(
        "AZURE_OPENAI_DEPLOYMENT_NAME",
        "AZURE_OPENAI_DEPLOYMENT"
    )

    api_version = get_env_value(
        "AZURE_OPENAI_API_VERSION",
        default="2024-02-15-preview"
    )

    if not deployment_name:
        raise Exception("AZURE_OPENAI_DEPLOYMENT_NAME 또는 AZURE_OPENAI_DEPLOYMENT 환경변수가 필요합니다.")

    client = AzureOpenAI(
        azure_endpoint=endpoint,
        api_key=api_key,
        api_version=api_version
    )

    system_prompt = """
너는 Slipwell 개인 매매동향 및 투자 행동 리포트를 작성하는 분석가다.

너의 역할:
- 입력된 주문, 체결, 시장, 행동, 성과 지표를 바탕으로 투자 행동을 해석한다.
- 단순 문장 변환이 아니라, 입력 지표를 근거로 한 구체적인 행동 개선 제안을 작성한다.
- 개선 제안은 특정 종목 매수/매도 추천이 아니라, 주문 방식과 체결 관리 습관 개선에 초점을 둔다.

규칙:
- 입력된 수치와 지표만 근거로 해석한다.
- 없는 내용을 추측하지 않는다.
- 투자 권유, 특정 종목 매수/매도 추천 표현은 쓰지 않는다.
- "사라", "팔아라", "매수 추천", "매도 추천" 같은 표현은 금지한다.
- 점수, 위험등급 표현은 쓰지 않는다.
- 과장된 표현은 피하고 보고서 문체로 작성한다.
- 반드시 JSON 형식으로만 응답한다.
- 마크다운 코드블록은 사용하지 않는다.
- 필드명은 diagnosis_text, prescription_text, full_report_text만 사용한다.
- 수치가 null이거나 "-"인 항목은 억지로 해석하지 않는다.
- LP 관련 지표는 원천 로그 확인용으로만 보고, 성과 판단이나 행동 개선 제안의 핵심 근거로 사용하지 않는다.

diagnosis_text 작성 기준:
- 주문 방식, 체결 결과, 시장 상태, 행동 유형을 종합적으로 진단한다.
- 평균 체결가, 체결 수량, 슬리피지율, 동조화 지수, 주문 불균형, 시장가 주문 비율, 매도 주문 비율, 스프레드, 호가 깊이 지표가 있으면 적극 반영한다.
- 보고서는 target_order_id의 주문 시점 기준 시장 데이터를 바탕으로 해석한다.
- 수치가 없는 항목은 언급하지 않는다.

prescription_text 작성 기준:
- 단순한 주의 문구가 아니라 행동 개선 제안으로 작성한다.
- 주문 방식 개선, 시장 상황 확인, 체결 가격 관리, 주문 규모 조절 관점에서 작성한다.
- 각 제안은 어떤 지표 때문에 필요한지 함께 설명한다.
- 시장가 주문 비율이 높으면 지정가 주문 활용, 주문 분할, 체결 가격 범위 확인을 제안한다.
- 매도 주문 비율이 높고 주문 불균형이 음수이면 매도 쏠림 구간에서 주문 전 시장 상태 확인을 제안한다.
- 슬리피지율이 불리하면 주문 전 호가, 스프레드, 체결 가능 가격 범위 확인을 제안한다.
- 스프레드가 크거나 호가 깊이가 부족하면 대량 시장가 주문보다 주문 분할과 가격 범위 확인을 제안한다.
- 동조화 지수가 높으면 군집 행동 또는 한 방향 쏠림 가능성을 고려해 즉시 주문보다 확인 절차를 강화하는 제안을 한다.

full_report_text 작성 기준:
- 종합 진단과 개선 제안을 자연스럽게 연결한 보고서 본문으로 작성한다.
- 성과 지표는 핵심 판단 기준이 아니라 참고 지표로만 다룬다.
"""

    user_prompt = f"""
아래 JSON은 가상 주식 시장 시뮬레이션에서 계산된 투자 행동 분석 데이터입니다.

응답 형식:
{{
  "diagnosis_text": "입력 지표를 근거로 한 투자 행동 종합 진단",
  "prescription_text": "입력 지표를 근거로 한 구체적인 행동 개선 제안",
  "full_report_text": "진단과 개선 제안을 포함한 전체 보고서 본문"
}}

작성 요구사항:
- diagnosis_text에는 사용자의 주문/체결/시장 상황을 종합 해석한다.
- prescription_text에는 사용자가 다음 시뮬레이션에서 개선할 수 있는 행동 제안을 작성한다.
- 개선 제안은 투자 권유가 아니라 주문 방식, 체결 관리, 시장 확인 습관에 대한 제안이어야 한다.
- 입력 데이터에 없는 내용을 만들지 않는다.
- 시장 데이터는 target_order_id의 주문 시점 기준으로 조회된 값이다.
- LP 관련 지표는 원천 로그 확인용으로만 보고, 성과 판단이나 행동 개선 제안의 핵심 근거로 사용하지 않는다.

입력 데이터:
{json.dumps(openai_input, ensure_ascii=False, indent=2)}
"""

    response = client.chat.completions.create(
        model=deployment_name,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.3
    )

    content = response.choices[0].message.content

    try:
        return safe_json_loads(content)
    except Exception:
        return {
            "diagnosis_text": content,
            "prescription_text": "",
            "full_report_text": content
        }


# ============================================================
# DOCX 생성
# ============================================================

def create_docx(openai_input, openai_output):
    font_name = "Malgun Gothic"

    doc = Document()
    apply_doc_font(doc, font_name)

    add_fonted_heading(doc, "[Slipwell AI] 개인 매매동향 및 투자 행동 리포트", level=1, font_name=font_name)

    add_fonted_heading(doc, "1. 주요 거래 지표 요약", level=2, font_name=font_name)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "항목", font_name)
    set_cell_text(hdr[1], "결과", font_name)

    report_context = openai_input["report_context"]
    order = openai_input["order_summary"]
    execution = openai_input["execution_summary"]
    market = openai_input["market_summary"]
    performance = openai_input["performance_summary"]
    behavior = openai_input["behavior_summary"]

    rows = [
        ("에이전트 ID", format_value(report_context.get("agent_id"))),
        ("전략 유형", format_value(report_context.get("strategy_type"))),
        ("종목명", format_value(report_context.get("symbol_name"))),
        ("분석 대상 주문 시각", format_value(report_context.get("target_order_timestamp"))),
        ("집행 주문", format_value(order.get("order_summary_text"))),
        ("주문 기준 현재가", format_value(execution.get("current_price"), "원")),
        ("현재가 기준 데이터", format_value(execution.get("current_price_source"))),
        ("평균 체결가", format_value(execution.get("average_execution_price"), "원")),
        ("체결 수량", format_value(execution.get("executed_quantity"), "주")),
        ("체결 금액", format_value(execution.get("trade_amount"), "원")),
        ("슬리피지율", format_value(execution.get("slippage_rate"), "%")),
        ("예상 기회손익", format_value(execution.get("opportunity_loss"), "원")),
        ("누적 시장가 주문 비율", format_value(order.get("market_order_ratio"))),
        ("누적 매도 주문 비율", format_value(order.get("sell_order_ratio"))),
        ("평균 체결 수량", format_value(order.get("avg_order_quantity"))),
        ("수익률", format_value(performance.get("return_rate"))),
        ("실현 손익", format_value(performance.get("realized_pnl"), "원")),
        ("미실현 손익", format_value(performance.get("unrealized_pnl"), "원")),
        ("총 자산 가치", format_value(performance.get("total_asset_value"), "원")),
        ("시장 지표 기준", format_value(market.get("interval_type"))),
        ("시장 지표 시각", format_value(market.get("feature_window_timestamp"))),
        ("동조화 지수", format_value(market.get("sync_index"))),
        ("주문 불균형", format_value(market.get("order_imbalance"))),
        ("구간 거래량", format_value(market.get("volume"))),
        ("구간 거래 횟수", format_value(market.get("trade_count"))),
        ("평균 스프레드", format_value(market.get("avg_spread"))),
        ("평균 매수호가 깊이", format_value(market.get("avg_bid_depth"))),
        ("평균 매도호가 깊이", format_value(market.get("avg_ask_depth"))),
        ("호가 깊이 불균형", format_value(market.get("depth_imbalance"))),
        ("LP 활성 비율", format_value(market.get("lp_active_ratio"))),
        ("행동 유형", format_value(behavior.get("behavior_type")))
    ]

    for item, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], item, font_name)
        set_cell_text(cells[1], value, font_name)

    add_fonted_heading(doc, "2. 종합 해석", level=2, font_name=font_name)
    add_fonted_paragraph(doc, openai_output.get("diagnosis_text", ""), font_name=font_name)

    add_fonted_heading(doc, "3. 개선 제안", level=2, font_name=font_name)
    add_fonted_paragraph(doc, openai_output.get("prescription_text", ""), font_name=font_name)

    add_fonted_heading(doc, "4. 전체 보고서", level=2, font_name=font_name)
    add_fonted_paragraph(doc, openai_output.get("full_report_text", ""), font_name=font_name)

    filename = f"Slipwell_개인매매동향리포트_{report_context['agent_id']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

    temp_dir = tempfile.gettempdir()
    local_path = os.path.join(temp_dir, filename)

    doc.save(local_path)

    return local_path, filename


# ============================================================
# Blob 업로드
# ============================================================

def upload_to_blob(local_path, filename):
    conn_str = os.environ["BLOB_CONNECTION_STRING"]
    container_name = os.environ["BLOB_CONTAINER_NAME"]

    blob_name = f"behavior/{filename}"

    blob_service_client = BlobServiceClient.from_connection_string(conn_str)
    blob_client = blob_service_client.get_blob_client(
        container=container_name,
        blob=blob_name
    )

    with open(local_path, "rb") as f:
        blob_client.upload_blob(f, overwrite=True)

    return blob_name, blob_client.url


# ============================================================
# report_output 저장
# ============================================================

def save_report_output(
    conn,
    agent_id,
    target_order_id,
    analysis_id,
    req_body,
    openai_input,
    openai_output,
    blob_name,
    blob_url
):
    cursor = conn.cursor()

    report_context = openai_input["report_context"]

    cursor.execute("""
        INSERT INTO dbo.report_output (
            agent_id,
            target_order_id,
            analysis_id,
            report_title,
            report_type,
            symbol_name,
            digital_twin_market_no,
            analysis_start_date,
            analysis_end_date,
            model_version,
            key_metrics_json,
            diagnosis_text,
            prescription_text,
            full_report_text,
            blob_name,
            blob_url
        )
        OUTPUT INSERTED.report_id
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """,
        agent_id,
        target_order_id,
        analysis_id,
        "[Slipwell AI] 개인 매매동향 및 투자 행동 리포트",
        report_context.get("report_type"),
        report_context.get("symbol_name"),
        report_context.get("digital_twin_market_no"),
        report_context.get("analysis_start_date"),
        report_context.get("analysis_end_date"),
        get_env_value("AZURE_OPENAI_DEPLOYMENT_NAME", "AZURE_OPENAI_DEPLOYMENT", default="gpt-4o-mini"),
        json.dumps(openai_input, ensure_ascii=False),
        openai_output.get("diagnosis_text"),
        openai_output.get("prescription_text"),
        openai_output.get("full_report_text"),
        blob_name,
        blob_url
    )

    report_id = cursor.fetchone()[0]
    conn.commit()
    return report_id


# ============================================================
# OpenAI 단독 테스트 API
# ============================================================

@app.route(route="test_openai", methods=["GET"])
def test_openai(req: func.HttpRequest) -> func.HttpResponse:
    try:
        test_input = {
            "report_context": {
                "agent_id": "test_agent",
                "strategy_type": "TEST",
                "symbol_name": "SAMSUNG",
                "digital_twin_market_no": "DTM_001",
                "analysis_start_date": "2026-06-26T09:00:00",
                "analysis_end_date": "2026-06-26T10:00:00",
                "report_type": "INVESTMENT_BEHAVIOR_DIAGNOSIS",
                "target_order_timestamp": "2026-06-26 09:00:00"
            },
            "performance_summary": {
                "return_rate": -0.05,
                "realized_pnl": -100000,
                "unrealized_pnl": -50000,
                "total_asset_value": 9850000
            },
            "order_summary": {
                "order_summary_text": "MARKET SELL / 100주",
                "market_order_ratio": 0.64,
                "sell_order_ratio": 0.72,
                "avg_order_quantity": 18.5
            },
            "execution_summary": {
                "current_price": 78500,
                "current_price_source": "market_ohlcv_1s",
                "average_execution_price": 76920,
                "executed_quantity": 100,
                "slippage_rate": -2.01,
                "opportunity_loss": 158000
            },
            "market_summary": {
                "interval_type": "1m",
                "current_price": 78500,
                "current_price_source": "market_ohlcv_1s",
                "feature_window_timestamp": "2026-06-26 09:00:00",
                "return_rate": -0.0201,
                "volatility": 0.018,
                "order_imbalance": -0.74,
                "sync_index": 0.74,
                "volume": 100,
                "trade_count": 1,
                "avg_spread": 500,
                "avg_bid_depth": 109208,
                "avg_ask_depth": 200434,
                "depth_imbalance": -0.2946,
                "lp_active_ratio": 0.35
            },
            "behavior_summary": {
                "behavior_type": "MARKET_SELL_BIAS",
                "behavior_summary_text": "시장가 주문 비율과 매도 주문 비율이 높게 나타난 행동 패턴입니다.",
                "main_causes": [
                    "시장가 주문 비율 증가",
                    "매도 주문 비중 증가",
                    "동조화 지수 상승"
                ]
            }
        }

        result = call_openai(test_input)

        return func.HttpResponse(
            json.dumps({
                "success": True,
                "openai_output": result
            }, ensure_ascii=False),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        return func.HttpResponse(
            json.dumps({
                "success": False,
                "message": str(e)
            }, ensure_ascii=False),
            status_code=500,
            mimetype="application/json"
        )


# ============================================================
# 보고서 생성 API
# ============================================================

@app.route(route="generate_report", methods=["POST"])
def generate_report(req: func.HttpRequest) -> func.HttpResponse:
    conn = None

    try:
        req_body = req.get_json()

        agent_id = req_body.get("agent_id")
        target_order_id = req_body.get("target_order_id")
        interval_type = req_body.get("interval_type", "1m")

        if not agent_id or not target_order_id:
            return func.HttpResponse(
                json.dumps({
                    "success": False,
                    "message": "agent_id와 target_order_id는 필수입니다."
                }, ensure_ascii=False),
                status_code=400,
                mimetype="application/json"
            )

        if interval_type not in ["1s", "1m", "1h", "1d"]:
            return func.HttpResponse(
                json.dumps({
                    "success": False,
                    "message": "interval_type은 1s, 1m, 1h, 1d 중 하나여야 합니다."
                }, ensure_ascii=False),
                status_code=400,
                mimetype="application/json"
            )

        conn = get_sql_connection()

        data = fetch_report_data(conn, agent_id, target_order_id, interval_type)
        openai_input = build_openai_input(agent_id, target_order_id, req_body, data)

        openai_output = call_openai(openai_input)

        local_path, filename = create_docx(openai_input, openai_output)
        blob_name, blob_url = upload_to_blob(local_path, filename)

        analysis_id = data["behavior"].analysis_id if data["behavior"] else None

        report_id = save_report_output(
            conn,
            agent_id,
            target_order_id,
            analysis_id,
            req_body,
            openai_input,
            openai_output,
            blob_name,
            blob_url
        )

        return func.HttpResponse(
            json.dumps({
                "success": True,
                "report_id": report_id,
                "blob_name": blob_name,
                "blob_url": blob_url,
                "message": "보고서 생성 완료"
            }, ensure_ascii=False),
            status_code=200,
            mimetype="application/json"
        )

    except Exception as e:
        return func.HttpResponse(
            json.dumps({
                "success": False,
                "message": str(e)
            }, ensure_ascii=False),
            status_code=500,
            mimetype="application/json"
        )

    finally:
        if conn:
            conn.close()